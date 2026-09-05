from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT=Path(__file__).resolve().parent/"Dokumentasi_TimManager_Lengkap.docx"
GREEN,BLUE,DARK,GRAY,INK="153D36","2E74B5","1F4D78","64748B","0F172A"

def font(run,size=11,color=INK,bold=False,italic=False):
 run.font.name="Calibri"; run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"),"Calibri"); run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"),"Calibri")
 run.font.size=Pt(size); run.font.color.rgb=RGBColor.from_string(color); run.bold=bold; run.italic=italic

def shade(cell,fill):
 p=cell._tc.get_or_add_tcPr(); s=p.find(qn("w:shd"))
 if s is None:s=OxmlElement("w:shd");p.append(s)
 s.set(qn("w:fill"),fill)

def geometry(table,widths):
 table.autofit=False; table.alignment=WD_TABLE_ALIGNMENT.LEFT; p=table._tbl.tblPr
 w=p.find(qn("w:tblW"))
 if w is None:w=OxmlElement("w:tblW");p.append(w)
 w.set(qn("w:w"),str(sum(widths)));w.set(qn("w:type"),"dxa")
 ind=p.find(qn("w:tblInd"))
 if ind is None:ind=OxmlElement("w:tblInd");p.append(ind)
 ind.set(qn("w:w"),"120");ind.set(qn("w:type"),"dxa")
 grid=table._tbl.tblGrid
 for c in list(grid):grid.remove(c)
 for width in widths:
  c=OxmlElement("w:gridCol");c.set(qn("w:w"),str(width));grid.append(c)
 for row in table.rows:
  for cell,width in zip(row.cells,widths):
   cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
   pr=cell._tc.get_or_add_tcPr();cw=pr.find(qn("w:tcW"))
   if cw is None:cw=OxmlElement("w:tcW");pr.append(cw)
   cw.set(qn("w:w"),str(width));cw.set(qn("w:type"),"dxa")

def table(doc,headers,rows,widths):
 t=doc.add_table(rows=1,cols=len(headers));t.style="Table Grid"
 header_props=t.rows[0]._tr.get_or_add_trPr();repeat=OxmlElement("w:tblHeader");repeat.set(qn("w:val"),"true");header_props.append(repeat)
 for i,h in enumerate(headers):
  shade(t.rows[0].cells[i],"E8EEF5");font(t.rows[0].cells[i].paragraphs[0].add_run(str(h)),9.5,DARK,True)
 for row in rows:
  cells=t.add_row().cells
  for i,v in enumerate(row):font(cells[i].paragraphs[0].add_run(str(v)),9.2)
 geometry(t,widths);doc.add_paragraph();return t

def bullets(doc,items,numbered=False):
 for index,item in enumerate(items,1):
  if numbered:
   p=doc.add_paragraph();p.paragraph_format.left_indent=Inches(.375);p.paragraph_format.first_line_indent=Inches(-.25)
   text=f"{index}. {item}"
  else:
   p=doc.add_paragraph(style="List Bullet");text=item
  p.paragraph_format.space_after=Pt(4);p.paragraph_format.line_spacing=1.25;font(p.add_run(text),10.5)

def remove_paragraph_borders(paragraph_or_style):
 ppr=paragraph_or_style._element.get_or_add_pPr()
 border=ppr.find(qn("w:pBdr"))
 if border is not None:ppr.remove(border)

def code_lines(doc,items):
 for item in items:
  p=doc.add_paragraph();p.paragraph_format.left_indent=Inches(.25);p.paragraph_format.space_after=Pt(2)
  run=p.add_run(item);run.font.name="Consolas";run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"),"Consolas");run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"),"Consolas");run.font.size=Pt(9);run.font.color.rgb=RGBColor.from_string(INK)

doc=Document();sec=doc.sections[0]
sec.page_width=Inches(8.5);sec.page_height=Inches(11);sec.top_margin=sec.bottom_margin=sec.left_margin=sec.right_margin=Inches(1);sec.header_distance=sec.footer_distance=Inches(.492)
normal=doc.styles["Normal"];normal.font.name="Calibri";normal._element.rPr.rFonts.set(qn("w:ascii"),"Calibri");normal._element.rPr.rFonts.set(qn("w:hAnsi"),"Calibri");normal.font.size=Pt(11);normal.paragraph_format.space_after=Pt(6);normal.paragraph_format.line_spacing=1.25
for name,size,color,before,after in [("Heading 1",16,INK,18,10),("Heading 2",13,INK,14,7),("Heading 3",12,INK,10,5)]:
 s=doc.styles[name];s.font.name="Calibri";s._element.rPr.rFonts.set(qn("w:ascii"),"Calibri");s._element.rPr.rFonts.set(qn("w:hAnsi"),"Calibri");s.font.size=Pt(size);s.font.bold=True;s.font.color.rgb=RGBColor.from_string(color);s.paragraph_format.space_before=Pt(before);s.paragraph_format.space_after=Pt(after);s.paragraph_format.keep_with_next=True
for name in ("List Bullet","List Number"):
 s=doc.styles[name];s.font.name="Calibri";s.font.size=Pt(10.5);s.paragraph_format.left_indent=Inches(.375);s.paragraph_format.first_line_indent=Inches(-.188);s.paragraph_format.space_after=Pt(4);s.paragraph_format.line_spacing=1.25
title_style=doc.styles["Title"];title_style.font.name="Calibri";title_style._element.rPr.rFonts.set(qn("w:ascii"),"Calibri");title_style._element.rPr.rFonts.set(qn("w:hAnsi"),"Calibri");title_style.font.size=Pt(30);title_style.font.bold=True;title_style.font.color.rgb=RGBColor.from_string("000000");remove_paragraph_borders(title_style)
h=sec.header.paragraphs[0];font(h.add_run("RUANGKERJA / TIMMANAGER  |  DOKUMENTASI SISTEM"),8.5,GRAY,True)
f=sec.footer.paragraphs[0];f.alignment=WD_ALIGN_PARAGRAPH.RIGHT;font(f.add_run("Dokumentasi versi 2.0  |  5 September 2026"),8.5,GRAY)

for _ in range(5):doc.add_paragraph()
p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;font(p.add_run("PANDUAN SISTEM DAN PENGGUNA"),10,"B7791F",True)
p=doc.add_paragraph(style="Title");remove_paragraph_borders(p);p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.paragraph_format.space_after=Pt(8);font(p.add_run("Dokumentasi RuangKerja"),30,"000000",True)
p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;font(p.add_run("Manajemen tim, proyek, tugas, file, notifikasi, dan laporan"),14,DARK)
doc.add_paragraph()
p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;font(p.add_run("Dokumentasi implementasi | Versi 2.0 | 5 September 2026"),10,GRAY,False,True)
doc.add_paragraph()
t=doc.add_table(rows=1,cols=1);geometry(t,[9360]);shade(t.cell(0,0),"EAF4F1");font(t.cell(0,0).paragraphs[0].add_run("Cakupan: autentikasi, profil, superadmin global, tim dan RBAC, proyek Kanban, kolaborasi tugas, file berversi, Fonnte global, Gmail, pencarian, dashboard, laporan, dan operasi produksi."),10.5,GREEN,True)
doc.add_page_break()

doc.add_heading("Daftar Isi",1)
bullets(doc,["Ringkasan sistem","Arsitektur dan sitemap","Peran dan hak akses","Model data dan tabel database","Alur kerja utama","Panduan penggunaan user","Panduan penggunaan superadmin","File dan keamanan upload","Integrasi Gmail dan Fonnte","Dashboard dan laporan","Deployment dan operasi","Indeks diagram Draw.io"],True)

doc.add_heading("Ringkasan Sistem",1)
doc.add_paragraph("RuangKerja adalah aplikasi kolaborasi tim berbasis Laravel. Unit kerja utama adalah Tim, Proyek/Papan, List Kanban, dan Tugas. Anggota berkolaborasi melalui assignee, checklist, komentar dengan mention, lampiran, chat grup, pengumuman, serta notifikasi.")
table(doc,["Area","Kemampuan utama"],[
["Akun","Registrasi, login, logout, profil, avatar, dan perubahan kata sandi."],
["Tim","Pembuatan tim, anggota, undangan email, dan role owner/admin/member."],
["Proyek","Papan Kanban, list status, ringkasan, chat, dan pengumuman."],
["Tugas","Prioritas, tenggat, drag-and-drop, multi-assignee, checklist, komentar, mention, aktivitas."],
["File","Upload privat, whitelist tipe/ukuran, preview, download, versioning, izin, dan antivirus opsional."],
["Integrasi","Gmail SMTP untuk undangan dan satu API key Fonnte global untuk notifikasi WhatsApp."],
["Analitik","Status tugas, beban anggota, aktivitas proyek, histori notifikasi, PDF dan XLSX."]],[1800,7560])

doc.add_heading("Arsitektur dan Sitemap",1)
doc.add_paragraph("Aplikasi menggunakan pola server-rendered Laravel dengan Blade dan Tailwind CSS. Route web dilindungi middleware auth, controller melakukan validasi dan otorisasi berbasis keanggotaan tim, Eloquent mengakses database, queue memproses notifikasi WhatsApp, dan storage local menyimpan file privat.")
bullets(doc,[
"Presentation: Blade views, komponen avatar/lampiran, JavaScript drag-and-drop, dan Tailwind CSS.",
"Application: controller, Form Request, action/service, mailable, job queue, dan scheduled command.",
"Domain/data: User, Team, Board, BoardList, Task, kolaborasi tugas, Attachment, dan log notifikasi.",
"Infrastructure: MySQL atau SQLite untuk test, storage privat, SMTP Gmail, Fonnte Send API, queue database, dan scheduler."])
doc.add_heading("Sitemap pengguna",2)
table(doc,["Halaman","Fungsi"],[
["/login dan /register","Masuk dan membuat akun."],["/","Beranda tim dan proyek."],["/search","Pencarian global proyek, tugas, dan file aktif."],["/reports","Dashboard laporan dan filter periode/proyek."],["/profile","Profil, avatar, kata sandi, dan pengaturan penerima WhatsApp."],["/admin","Dashboard khusus superadmin untuk statistik, pengguna, dan audit akses."],["/teams/{team}","Pengaturan tim, anggota, role, dan undangan."],["/boards/{board}","Kanban dan filter tugas."],["/boards/{board}/summary","Ringkasan proyek."],["/boards/{board}/chat","Chat grup dan lampiran."],["/boards/{board}/announcements","Pengumuman dan lampiran."],["/boards/{board}/tasks/{task}","Detail, assignee, checklist, komentar, aktivitas, dan file."],["/attachments/{attachment}/preview","Preview dan riwayat versi file."]],[3000,6360])

doc.add_heading("Peran dan Hak Akses",1)
doc.add_heading("Role global",2)
table(doc,["Role","Akses"],[
["User","Mengakses tim, proyek, tugas, file, laporan, dan profil sesuai keanggotaan."],
["Superadmin","Memiliki akses tambahan ke /admin untuk statistik platform, pengelolaan role global, status akun, dan audit."]],[2200,7160])
doc.add_paragraph("Role global tidak menggantikan role per tim. Seorang superadmin tetap mengikuti aturan owner, admin, atau member ketika mengelola resource tim.")
doc.add_heading("Role per tim",2)
table(doc,["Aksi","Owner","Admin","Member"],[
["Melihat tim/proyek yang diikuti","Ya","Ya","Ya"],["Mengubah/hapus tim","Ya","Tidak","Tidak"],["Kelola anggota","Ya","Ya","Tidak"],["Mengangkat admin","Ya","Tidak","Tidak"],["Buat/ubah/hapus proyek","Ya","Ya","Tidak"],["Kelola tugas dan kolaborasi","Ya","Ya","Ya"],["Hapus komentar orang lain","Ya","Ya","Tidak"],["Atur izin download proyek","Ya","Ya","Tidak"],["Melihat laporan yang dapat diakses","Ya","Ya","Ya"]],[4110,1750,1750,1750])
doc.add_paragraph("Resource lintas tim umumnya dikembalikan sebagai HTTP 404 agar keberadaan data tidak terungkap. ID assignee dan mention diverifikasi sebagai anggota tim. Akun dengan is_active bernilai false tidak dapat login. Sistem melindungi superadmin aktif terakhir agar tidak diturunkan atau dinonaktifkan.")

doc.add_heading("Model Data dan Tabel Database",1)
tables=[("users","Akun, profil, avatar, global_role, is_active, dan kredensial."),("admin_audit_logs","Actor, target user, action, before/after, IP, dan user agent."),("teams","Ruang kolaborasi; owner_id ke users."),("team_user","Keanggotaan dan role admin/member."),("team_invitations","Undangan email bertoken."),("boards","Proyek dan download_permission."),("board_lists","Kolom/status Kanban."),("tasks","Data tugas dan creator."),("task_user","Pivot multi-assignee."),("task_checklist_items","Checklist dan status selesai."),("task_comments","Komentar tugas."),("task_comment_mentions","Pivot mention komentar."),("task_activities","Audit aktivitas dan metadata JSON."),("board_messages","Chat grup proyek."),("announcements","Pengumuman dan pin."),("attachments","File polymorphic, versi, scan, file aktif."),("whatsapp_connections","Nomor penerima, consent, preferensi, timezone, dan quiet hours."),("whatsapp_notification_logs","Status pengiriman WhatsApp."),("email_notification_logs","Status pengiriman email."),("jobs / failed_jobs / job_batches","Infrastruktur queue."),("sessions / cache / cache_locks","Sesi dan cache.")]
table(doc,["Tabel","Tujuan"],tables,[3000,6360])
doc.add_heading("Relasi kunci",2)
bullets(doc,[
"User memiliki Team sebagai owner dan bergabung melalui team_user.",
"Team memiliki banyak Board; Board memiliki BoardList; BoardList memiliki Task.",
"Task memiliki assignee, checklist, komentar, aktivitas, dan attachment.",
"TaskComment memiliki mention melalui task_comment_mentions.",
"Attachment polymorphic untuk Task, BoardMessage, dan Announcement; root_attachment_id membentuk keluarga versi.",
"User memiliki satu WhatsappConnection; connection memiliki banyak WhatsappNotificationLog.",
"User dapat menjadi actor atau target pada AdminAuditLog. API key Fonnte tidak disimpan pada record user."])
doc.add_heading("Integritas dan penghapusan",2)
bullets(doc,[
"Agregat tim, proyek, list, tugas, dan data turunan memakai cascade sesuai kepemilikan.",
"Actor activity dan root attachment dapat menjadi null untuk mempertahankan histori.",
"File fisik dihapus ketika record attachment dihapus.",
"Versi baru menonaktifkan versi sebelumnya di dalam transaksi."])

doc.add_heading("Alur Kerja Utama",1)
doc.add_heading("Membuat dan mengelola tugas",2)
bullets(doc,[
"Anggota membuka proyek yang dapat diakses.",
"Mengisi list, judul, deskripsi, prioritas, tenggat, multi-assignee, dan lampiran.",
"Server memvalidasi board, list, anggota, field, dan file.",
"Task, pivot assignee, attachment, dan activity created disimpan dalam transaksi.",
"Notifikasi WhatsApp dibuat sesuai koneksi, preferensi, dan quiet hours.",
"Kartu dapat dipindahkan antar-list; detail tugas memuat checklist, komentar, mention, file, dan activity log."],True)
doc.add_heading("Mengunggah versi file",2)
bullets(doc,[
"Validasi jumlah, ukuran maksimum 10 MB, MIME, dan ekstensi whitelist.",
"Jika antivirus aktif, file dipindai sebelum masuk storage.",
"File disimpan pada disk local privat beserta metadata scan.",
"Versi baru wajib memiliki ekstensi yang sama.",
"Versi sebelumnya menjadi tidak aktif dan versi terbaru menjadi is_current.",
"Preview/download memeriksa membership dan download_permission."],True)
doc.add_heading("Notifikasi",2)
bullets(doc,[
"Event aplikasi membentuk pesan notifikasi.",
"Preferensi dan status koneksi diperiksa.",
"Quiet hours menentukan waktu pengiriman.",
"Log pending dibuat dengan idempotency key; job masuk queue.",
"Job memakai FONNTE_API_KEY global untuk memanggil Fonnte Send API dan mencatat sent, failed, atau skipped.",
"Email undangan dikirim via SMTP dan dicatat di email_notification_logs."],True)

doc.add_heading("Panduan Penggunaan User",1)
doc.add_heading("Memulai",2)
bullets(doc,["Registrasi lalu login.","Lengkapi profil, jabatan, telepon, bio, avatar, dan kata sandi.","Buat tim; pembuat otomatis menjadi Owner.","Undang anggota lewat Gmail atau tambahkan akun terdaftar.","Buat proyek pada tim yang dikelola."],True)
doc.add_heading("Kanban dan tugas",2)
bullets(doc,["Buka proyek dan pilih Buat tugas.","Pilih satu atau beberapa penanggung jawab.","Gunakan filter teks, anggota, prioritas, dan tenggat.","Seret kartu untuk memperbarui status.","Klik judul untuk membuka detail.","Kelola checklist, komentar, mention, dan file."],True)
doc.add_heading("Pencarian dan laporan",2)
bullets(doc,["Gunakan pencarian pada Beranda atau header proyek.","Hasil hanya menampilkan resource yang berhak diakses.","Buka Laporan untuk status tugas, workload, aktivitas, dan histori notifikasi.","Pilih proyek dan periode 7/30/90 hari.","Unduh PDF atau Excel."],True)

doc.add_heading("Notifikasi WhatsApp",2)
bullets(doc,["Buka Profil lalu pilih Koneksi WhatsApp.","Masukkan nomor penerima dengan kode negara, misalnya 628123456789.","Pilih notifikasi tugas, perubahan tugas, chat, pengumuman, dan pengingat tenggat yang diperlukan.","Atur zona waktu dan jam tenang.","Centang persetujuan, aktifkan pengiriman, lalu simpan.","Pilih Kirim pesan uji. Jika gagal, hubungi administrator karena API key Fonnte dikelola secara global."],True)

doc.add_page_break()
doc.add_heading("Panduan Penggunaan Superadmin",1)
doc.add_paragraph("Dashboard superadmin tersedia hanya untuk akun dengan global_role superadmin dan is_active bernilai true.")
doc.add_heading("Mengaktifkan superadmin pertama",2)
code_lines(doc,["/opt/alt/php84/usr/bin/php artisan users:promote-superadmin email@domain.com --force"])
doc.add_paragraph("Ganti email dengan akun yang sudah terdaftar. Login ulang setelah perintah berhasil, lalu buka /admin.")
doc.add_heading("Menggunakan dashboard",2)
bullets(doc,["Gunakan tombol panah untuk memperkecil atau memperbesar sidebar desktop. Pada ponsel, tombol menu membuka drawer.","Periksa statistik pengguna, pengguna aktif, superadmin, tim, proyek, dan tugas.","Cari pengguna berdasarkan nama atau email. Filter hasil berdasarkan role global dan status.","Pilih User atau Superadmin, tentukan status aktif, lalu pilih Simpan.","Periksa Aktivitas superadmin untuk melihat actor, target, tindakan, dan waktu perubahan."],True)
doc.add_heading("Pengamanan akun",2)
bullets(doc,["Akun nonaktif tidak dapat login dan sesi aktifnya dihapus.","Setiap perubahan role atau status dicatat dalam admin_audit_logs.","Sistem menolak tindakan yang akan menghilangkan superadmin aktif terakhir.","Gunakan akun superadmin hanya untuk administrasi platform. Gunakan role tim untuk pekerjaan proyek sehari-hari."])

doc.add_heading("File dan Keamanan Upload",1)
table(doc,["Kontrol","Implementasi"],[
["Jumlah","Maksimal 5 file per request."],["Ukuran","Maksimal 10 MB per file."],["Tipe","PDF, Office, TXT/CSV, JPEG/PNG/WebP, dan ZIP."],["Penyimpanan","Disk local privat melalui controller terotorisasi."],["Preview","Inline untuk image/PDF/text; tipe lain melalui download."],["Versioning","root_attachment_id, version, is_current, dan histori uploader."],["Antivirus","ClamAV opsional; fail-closed tersedia."],["Izin proyek","team, managers, atau uploader/owner."],["Header aman","nosniff dan CSP untuk respons inline."]],[2300,7060])
doc.add_paragraph("Produksi disarankan memasang ClamAV, mengaktifkan scanner, dan mempertahankan fail-closed.")

doc.add_heading("Integrasi Gmail dan Fonnte",1)
doc.add_heading("Gmail SMTP",2)
table(doc,["Variabel","Nilai contoh"],[["MAIL_MAILER","smtp"],["MAIL_SCHEME","smtp"],["MAIL_HOST","smtp.gmail.com"],["MAIL_PORT","587"],["MAIL_USERNAME","alamat Gmail aktif"],["MAIL_PASSWORD","Google App Password tanpa spasi"],["MAIL_FROM_ADDRESS","alamat pengirim"]],[3000,6360])
doc.add_paragraph("Aktifkan Verifikasi 2 Langkah, buat Google App Password, isi .env, lalu bersihkan cache konfigurasi. Jangan commit kredensial atau menampilkannya pada tangkapan layar.")
doc.add_heading("Fonnte Send API",2)
bullets(doc,[
"Administrator menyimpan satu token perangkat pada FONNTE_API_KEY di .env. Token tidak dimasukkan oleh user dan tidak disimpan pada whatsapp_connections.",
"Setiap user menyimpan nomor penerima, persetujuan, preferensi event, timezone, dan jam tenang.",
"Test connection dibatasi throttle.",
"Queue worker wajib aktif untuk pengiriman asinkron.",
"Scheduler menjalankan reminder tenggat; quiet hours memakai timezone pengguna.",
"Riwayat status tersedia pada profil dan laporan."])

doc.add_heading("Dashboard dan Laporan",1)
bullets(doc,[
"Metrik: total, selesai, dikerjakan/aktif, dan terlambat.",
"Beban kerja: total, aktif, selesai, dan terlambat per assignee.",
"Aktivitas per proyek dalam periode 7, 30, atau 90 hari.",
"Riwayat WhatsApp pengguna dan riwayat email tim yang dapat diakses.",
"Ekspor PDF dan workbook XLSX tanpa dependency eksternal tambahan."])
doc.add_paragraph("Definisi selesai menggunakan list bernama Selesai dan batal menggunakan list Batal. Pertahankan nama tersebut untuk konsistensi laporan.")

doc.add_heading("Deployment dan Operasi",1)
doc.add_heading("Persyaratan produksi",2)
bullets(doc,["PHP minimal 8.4.1. Pada hosting saat ini gunakan /opt/alt/php84/usr/bin/php.","Document root mengarah ke public atau aplikasi menyediakan public/index.php sesuai konfigurasi hosting.","Direktori storage dan bootstrap/cache dapat ditulis oleh user web hosting.","APP_ENV=production, APP_DEBUG=false, APP_URL memakai domain HTTPS, serta APP_KEY tetap dan tidak diganti setelah produksi.","QUEUE_CONNECTION=database, database sudah dibuat, dan public/storage terhubung ke storage/app/public."])
doc.add_heading("Pembaruan aplikasi",2)
code_lines(doc,[
"cd /home/u930607946/domains/ruangkerja.sivmi.id/public_html",
"git pull origin main",
"/opt/alt/php84/usr/bin/php artisan optimize:clear",
"/opt/alt/php84/usr/bin/php artisan migrate --force",
"/opt/alt/php84/usr/bin/php artisan storage:link",
"/opt/alt/php84/usr/bin/php artisan optimize"])
doc.add_paragraph("Jika server tidak memiliki npm, jalankan npm run build di komputer pengembang lalu unggah folder public/build. Pastikan manifest.json dan assets berada langsung di public/build, bukan public/build/build.")
doc.add_heading("Konfigurasi Fonnte global",2)
code_lines(doc,["FONNTE_API_KEY=token_perangkat_fonnte","FONNTE_API_URL=https://api.fonnte.com/send"])
doc.add_paragraph("Setelah .env diubah, jalankan optimize:clear lalu optimize. Verifikasi konfigurasi tanpa mencetak token dengan perintah tinker yang memeriksa apakah nilai services.fonnte.api_key terisi.")
doc.add_heading("Queue dan scheduler",2)
doc.add_paragraph("Buat dua cron job terpisah dengan jadwal setiap menit. Gunakan path PHP 8.4 agar Composer tidak dijalankan oleh PHP 8.3.")
code_lines(doc,[
"/opt/alt/php84/usr/bin/php /home/u930607946/domains/ruangkerja.sivmi.id/public_html/artisan queue:work database --stop-when-empty --sleep=1 --tries=3 --timeout=60",
"/opt/alt/php84/usr/bin/php /home/u930607946/domains/ruangkerja.sivmi.id/public_html/artisan schedule:run"])
doc.add_paragraph("Pada panel hosting yang tidak menyediakan daemon, pola stop-when-empty cocok dijalankan setiap menit. Log cron dapat diarahkan ke storage/logs bila panel mendukung output.")
doc.add_heading("Antivirus",2)
table(doc,["Variabel","Keterangan"],[["ATTACHMENT_ANTIVIRUS_ENABLED","true untuk mengaktifkan ClamAV."],["ATTACHMENT_ANTIVIRUS_COMMAND","Lokasi/perintah clamscan."],["ATTACHMENT_ANTIVIRUS_TIMEOUT","Batas pemindaian dalam detik."],["ATTACHMENT_ANTIVIRUS_FAIL_CLOSED","Menolak upload ketika scanner gagal."]],[3600,5760])
doc.add_heading("Quality gate",2)
bullets(doc,["php artisan test --compact","vendor/bin/pint --format agent","php artisan view:cache","npm run build","php artisan migrate:status"])
doc.add_paragraph("Status saat dokumentasi dibuat: 107 test lulus dengan 425 assertions; Blade cache dan build frontend berhasil.")

doc.add_heading("Indeks Diagram Draw.io",1)
table(doc,["Halaman","Isi"],[
["01 - Sitemap","Navigasi publik, workspace, tim, proyek, tugas, file, laporan, dan superadmin."],["02 - ERD Core","User, akses global, admin audit, Team, membership, Board, List, dan Task."],["03 - ERD Collaboration & File","Assignee, checklist, comment, mention, activity, attachment, dan notifikasi."],["04 - RBAC","Superadmin global serta Owner, Admin, dan Member pada tim."],["05 - Database Catalog","Inventaris tabel domain dan infrastruktur."],["06 - Task Flow","Pembuatan dan pembaruan tugas serta kolaborasi."],["07 - File Flow","Validasi, antivirus, storage, preview, versioning, dan permission."],["08 - Notification & Reports","Gmail, Fonnte global, queue, log, dashboard, dan ekspor."],["09 - User Guide","Langkah penggunaan user dari registrasi sampai laporan."],["10 - Deployment","PHP 8.4, .env, queue, scheduler, storage, dan quality gate."],["11 - Superadmin Guide","Promosi akun, dashboard, pengelolaan user, dan audit."]],[2700,6660])
doc.add_paragraph("Berkas sumber diagram: docs/timmanager-documentation.drawio. Buka dengan diagrams.net dan gunakan tab halaman di bagian bawah.")

doc.core_properties.title="Dokumentasi Lengkap RuangKerja";doc.core_properties.subject="Sitemap, ERD, RBAC, database, flowchart, integrasi, deployment, serta panduan user dan superadmin";doc.core_properties.author="RuangKerja";doc.core_properties.keywords="Laravel, Kanban, RBAC, ERD, Fonnte, WhatsApp, Gmail, Superadmin"
doc.save(OUT);print(OUT)
